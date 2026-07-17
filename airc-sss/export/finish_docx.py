#!/usr/bin/env python3
# LOCKED docx cover + footer post-process. Called by build.sh.
# Adds the cover (title once, subtitle, title-less journey graphic, two-line
# byline) BEFORE the TOC, then centered footer page numbers with the cover page
# unnumbered. Do not drift.
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

infile, graphic = sys.argv[1], sys.argv[2]
doc = Document(infile)
body = doc.element.body

def carlito(run, size, bold=False):
    run.font.name = 'Carlito'
    r = run._element.rPr.rFonts if run._element.rPr is not None and run._element.rPr.rFonts is not None else None
    rpr = run._element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    for a in ('w:ascii','w:hAnsi','w:cs'): rf.set(qn(a),'Carlito')
    run.font.size = Pt(size); run.bold = bold

# Build cover paragraphs (appended, then moved to the very front in order)
p_title = doc.add_paragraph(); p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
carlito(p_title.add_run('Student Journey Gap Analysis'), 26, True)

p_sub = doc.add_paragraph(); p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
carlito(p_sub.add_run('Student Support and Success Committee · AI Resource Center, Domain 5'), 13)

p_img = doc.add_paragraph(); p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_img.add_run().add_picture(graphic, width=Inches(6.3))

p_by = doc.add_paragraph(); p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p_by.add_run('Prepared by Michelle Blomberg, Co-chair,'); carlito(r1, 13)
r1.add_break()
carlito(p_by.add_run('Student Support and Success domain, AI Resource Center.'), 13)

p_pb = doc.add_paragraph(); p_pb.add_run().add_break(WD_BREAK.PAGE)

# Move to front in the order: title, subtitle, image, byline, pagebreak
for p in (p_pb, p_by, p_img, p_sub, p_title):
    body.remove(p._p); body.insert(0, p._p)

# Footer: centered PAGE field on primary footer; cover (first page) unnumbered
sec = doc.sections[0]
sec.different_first_page_header_footer = True
# empty first-page footer already (cover -> no number)
fp = sec.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run()
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), ' PAGE ')
t = OxmlElement('w:r'); tt = OxmlElement('w:t'); tt.text = '1'; t.append(tt); fld.append(t)
run._element.addnext(fld)

# Give every table visible single-line borders so they do not float.
def _set_borders(table):
    tblPr = table._tbl.tblPr
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    b = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement('w:'+edge)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4')
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),'BFB2C4')
        b.append(e)
    tblPr.append(b)
for _t in doc.tables:
    _set_borders(_t)

doc.save(infile)
print('docx cover + footer + table borders applied')
